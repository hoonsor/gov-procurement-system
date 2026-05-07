#!/usr/bin/env python3
"""
投標書範本產生器
自動生成政府採購投標文件
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
from pathlib import Path

def create_bid_document(tender_info, company_info):
    """產生投標書 Word 文件"""
    doc = Document()
    
    # 標題
    title = doc.add_heading('投  標  書', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 投標日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f"中華民國 {datetime.now().year} 年 {datetime.now().month} 月 {datetime.now().day} 日")
    date_run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # 一、投標廠商資料
    doc.add_heading('一、投標廠商資料', level=1)
    
    table1 = doc.add_table(rows=4, cols=2)
    table1.style = 'Table Grid'
    
    data1 = [
        ("公司名稱", company_info.get("name", "")),
        ("統一編號", company_info.get("tax_id", "")),
        ("地址", company_info.get("address", "")),
        ("聯絡人/電話", f"{company_info.get('contact', '')} / {company_info.get('phone', '')}")
    ]
    
    for i, (label, value) in enumerate(data1):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # 二、投標標的
    doc.add_heading('二、投標標的', level=1)
    
    table2 = doc.add_table(rows=5, cols=2)
    table2.style = 'Table Grid'
    
    data2 = [
        ("採購案號", tender_info.get("tender_id", "")),
        ("採購名稱", tender_info.get("title", "")),
        ("預算金額", f"新臺幣 {tender_info.get('budget', 0):,} 元"),
        ("截止投標日期", tender_info.get("deadline", "")),
        ("履約地點", tender_info.get("location", "機關指定地點"))
    ]
    
    for i, (label, value) in enumerate(data2):
        table2.rows[i].cells[0].text = label
        table2.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # 三、投標報價
    doc.add_heading('三、投標報價', level=1)
    
    table3 = doc.add_table(rows=4, cols=2)
    table3.style = 'Table Grid'
    
    total = company_info.get("bid_amount", 0)
    
    data3 = [
        ("投標總價", f"新臺幣 {total:,} 元整"),
        ("投標單價", f"詳見估價單"),
        ("投標有效期", f"自投標日起 {company_info.get('validity_days', 90)} 天內有效"),
        ("付款方式", company_info.get("payment_terms", "依政府採購法規定"))
    ]
    
    for i, (label, value) in enumerate(data3):
        table3.rows[i].cells[0].text = label
        table3.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # 四、投標聲明
    doc.add_heading('四、投標聲明', level=1)
    
    declarations = [
        "1. 本投標廠商願遵守政府採購法及相關規定辦理。",
        "2. 本投標廠商所投標的均為全新品，符合採購規格要求。",
        "3. 本投標廠商願提供所需之相關證明文件，供機關審查。",
        "4. 本投標廠商如得標，將依約定時間完成履約。",
        "5. 本投標廠商絕不以任何形式影響招標之公平性。",
        "6. 本投標廠商最近三年內無重大違法紀錄。"
    ]
    
    for dec in declarations:
        doc.add_paragraph(dec)
    
    doc.add_paragraph()
    
    # 五、附加文件
    doc.add_heading('五、附加文件', level=1)
    
    docs = [
        "□ 公司登記證明文件",
        "□ 統一編號編配證明",
        "□ 最近三年納稅證明",
        "□ 實績證明文件",
        "□ 品質認證證書",
        "□ 其他相關文件"
    ]
    
    for d in docs:
        doc.add_paragraph(d)
    
    doc.add_paragraph()
    
    # 六、切結事項
    doc.add_heading('六、切結事項', level=1)
    
    doc.add_paragraph("本廠商茲聲明以上所填資料均屬事實，如經查證不實，願負一切法律責任。")
    
    doc.add_paragraph()
    
    # 簽章欄
    doc.add_paragraph("此致")
    doc.add_paragraph(f"中華民國教育部{ tender_info.get('department', '') }")
    
    doc.add_paragraph()
    
    sign_para = doc.add_paragraph()
    sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_para.add_run("投標廠商（加蓋印章）：________________")
    
    doc.add_paragraph()
    
    sign_para2 = doc.add_paragraph()
    sign_para2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_para2.add_run("負責人（簽章）：________________")
    
    return doc

def save_bid_document(tender_info, company_info, filepath=None):
    """儲存投標書"""
    if filepath is None:
        filepath = Path(__file__).parent.parent / "templates" / f"投標書_{tender_info.get('tender_id', 'unknown')}.docx"
    
    filepath = Path(filepath)
    filepath.parent.mkdir(parents=True, exist_ok=True)
    
    doc = create_bid_document(tender_info, company_info)
    doc.save(str(filepath))
    
    print(f"✅ 投標書已產生: {filepath}")
    return filepath

# 預設公司資料
DEFAULT_COMPANY = {
    "name": "冠宇資訊科技有限公司",
    "tax_id": "12345678",
    "address": "台北市大安區復興南路一段100號12樓",
    "contact": "陳大明",
    "phone": "02-2771-5566",
    "bid_amount": 2350000,
    "validity_days": 90,
    "payment_terms": "依政府採購法規定，依完工進度分期付款"
}

DEFAULT_TENDER = {
    "tender_id": "TP-2026-001",
    "title": "校園網路設備更新採購案",
    "budget": 2500000,
    "deadline": "2026-05-25",
    "department": "資訊中心",
    "location": "台北市大安區和平東路一段187號"
}

if __name__ == "__main__":
    from pathlib import Path
    
    templates_dir = Path(__file__).parent.parent / "templates"
    templates_dir.mkdir(parents=True, exist_ok=True)
    
    doc = save_bid_document(DEFAULT_TENDER, DEFAULT_COMPANY)
    print(f"📄 投標書路徑: {doc}")